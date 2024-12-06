# pip install 'fastapi[standard]'
# pip install docxtpl
# pip install docx2pdf
# pip install python-docx

# Used for API calls
from fastapi import FastAPI, Form, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse

# Used in conversion of required data, when converting from UploadFile to other type
import json
import shutil

# Used for inserting data in template
from docxtpl import DocxTemplate
from docx import Document
from http import HTTPStatus

# Used for converting docx-format to pdf-format
import os
from docx2pdf import convert

# Used for unit testing
import unittest
from copy import deepcopy

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)

@app.post('/insert-dynamic-data/')
async def insert_dynamic_data(
        templateFile: UploadFile = File(...),
        pdfName: str = Form(...),
        contextFile: UploadFile = File(...),
        images: list[UploadFile] = File(None),
    ):
    # Extracts the context-dictionary from .json file
    context = json.loads(contextFile.file.read())
    
    # Save a copy pf the uploaded template file
    templatePath = 'uploadedTemplate.docx'
    with open(templatePath, "wb") as buffer:
        shutil.copyfileobj(templateFile.file, buffer)
    
    tpl = DocxTemplate(templatePath)

    context = InitializeContext(context)
    
    errMsg, valid = ValidateVariables(templatePath, context)
    if valid == False:
        return ValueError, errMsg

    # Inserts data from the context-dictionary to the template
    tpl.render(context)
    
    tpl.save(f'{pdfName}.docx')
    InsertPageNumbers(f'{pdfName}.docx')
    return ConvertDocxToPDF(pdfName)

def ConvertDocxToPDF(path: str):
    convert(f'{path}.docx', f'{path}.pdf')
    os.remove(f'{path}.docx')

    return True, FileResponse(f'{path}.pdf', media_type = 'application/pdf', filename = path)

def ValidateVariables(path: str, context: dict):
    doc = Document(path)
    valid = True

    keysContained = []
    keysNotContained = []
    values = []
    valuesNotInputted = []
    valuesInputted = []

    for p in doc.paragraphs:
        # All variables must begin with "{{", therefore we can find them simply by looking for this
        if p.text.__contains__('{{') and not (p.text.__contains__('{% for') or p.text.__contains__('{% endfor')):
            tempValues = p.text.rsplit('{{')

            for temp in tempValues:
                if not temp.__contains__('}}'):
                    tempValues.remove(temp)

            # I don't know why, but it only seems to be able to do one of the variables per for-loop...
            # It doesn't work correctly with a single for-loop, there must be one for each variable in the tempValues list
            for _ in range(len(tempValues)):
                for tempValue in tempValues:
                    tempValues.remove(tempValue)
                    tempValue = tempValue.split('}}')[0]
                    tempValues.append(tempValue)

            for temp in tempValues:
                if not values.__contains__(temp):
                    values.append(temp)

    # Ignores lists, because these are not acounted for as variables in the template
    # It is also not seen as an error, if these are empty
    for key in context:
        if not isinstance(context[key], list): keysNotContained.append(key)
    for value in values:
        valuesNotInputted.append(value)
    
    for value in values:
        for key in context:
            if value == key:
                keysContained.append(key)
                keysNotContained.remove(key)

                valuesInputted.append(value)
                valuesNotInputted.remove(value)
    
    # Builds the exception string, so it can be returned
    errorMsg = ''
    if not len(keysNotContained) == 0:
        for key in keysNotContained:
            errorMsg += f'{key} was not found in template. '
        valid = False
    
    if not len(valuesNotInputted) == 0:
        for value in valuesNotInputted:
            errorMsg += f'{value} was not found in context. '
        valid = False
    
    return errorMsg, valid

# Inserts the total price to the context dictionary, this is done to automate the proccess
# I would suggest moving this out of the python file and place it in the WebApp
def InsertTotalPrices(context: dict):
    index = 0
    totalPrice = 0
    for i in context['itemsTable']:
        context['itemsTable'][index]['priceTotal'] = i['units'] * i['priceUnit']
        totalPrice += context['itemsTable'][index]['priceTotal']

        index += 1

    context['itemsTable'].append(
            {
                'type' : 'Pris eks. moms',
                'units' : '',
                'priceUnit' : '',
                'priceTotal' : "%.2f" % totalPrice,
            }
    )
    context['itemsTable'].append(
            {
                'type' : 'Pris inkl. moms',
                'units' : '',
                'priceUnit' : '',
                'priceTotal' : "%.2f" % (totalPrice * 1.25),
            }
    )
    context['Total_Price'] = "%.2f" % (totalPrice * 1.25)
    return context

def InsertPageNumbers(path: str):
    doc = Document(path)

    pageNumber = 1
    for p in doc.paragraphs:
        if p.text.__contains__('{{'):
            subStr = p.text
            subStr = subStr.split('{{Current_Page}}')
            subStr = subStr[0]+str(pageNumber)+subStr[1]
            p.text = subStr
            pageNumber += 1
    
    doc.save(path)
    tpl = DocxTemplate(path)

    context = {'Total_Pages' : str(pageNumber - 1)}
    tpl.render(context)

    tpl.save(path)

# Completes the context dictionary, by adding the last required variables
# I would suggest moving this out of the python file and place it in the WebApp
def InitializeContext(context):
    context = InsertTotalPrices(context)
    context['Current_Page'] = '{{Current_Page}}'
    context['Total_Pages'] = '{{Total_Pages}}'
    return context

# ------------------------------------------------------------------------------------------- #
# ----------------------------------------Unit Tests:---------------------------------------- #


contextDict = {
    "Name" : "Magnus Næhr",
    "Adress" : "Mølletoften 1",
    "City_postcode" : "Lyngby, 2800",
    "Customer_Number" : 123456,
    "Order_Number" : 123456,
    "Current_Date" : "05-12-2024",
    "Invoice_Type" : "Ordrebekræftelse",
    "Offer_Name" : "Dette er tilbudet",
    "Date_of_Execution" : "13-12-2024",
    "Floor_Elevator" : "2. etage, m. elevator.",
    "Square_Meters" : 28,
    "Parking" : "Perkeringsplads ude foran bolig.",
    "Time_Estimation" : "2 timer.",
    "Task_Description" : "Dette er opgave beskrivelsen.",
    "Second_Adresses" : [
        "Dette er en sekundær addresse.",
        "Dette er også en sekundær adresse."
    ],
    "Comments" : [
        "Dette er en kommentar",
        "Dette er også en kommentar"
    ],
    "Agreed_Date" : "13-12-2024",
    "Start_Time" : "10:30 - 11:30",
    "Agreed_Date_Equipment" : "14-12-2024",
    "Start_Time_Equipment" : "11:00 - 12:00",
    "Customer_Phone" : "12345678",
    "Total_Price" : 0,
    "Reg_Number" : 123456,
    "Account_Number" : 123456,
    "Order_Id" : 123456,
    "Details" : [
        "Dette er en detalje.",
        "Dette er også en detalje"
    ],
    "itemsTable" : [
        {
            "type" : "Fast pris for flytning",
            "units" : 1,
            "priceUnit" : 9999.95,
            "priceTotal" : 0
        },
        {
            "type" : "Fast pris for nedpakning",
            "units" : 1,
            "priceUnit" : 249.95,
            "priceTotal" : 0
        },
        {
            "type" : "Fast pris for opbevaring",
            "units" : 1,
            "priceUnit" : 349.95,
            "priceTotal" : 0
        },
        {
            "type" : "Pris for udlejning af udstyr",
            "units" : 8,
            "priceUnit" : 199.95,
            "priceTotal" : 0
        },
        {
            "type" : "Tungløft",
            "units" : 7,
            "priceUnit" : 99.95,
            "priceTotal" : 0
        }
    ]
}

testTemplatePath = 'docxtpl/HC Andersen Flyttefirma Template.docx'
testPDFName = 'HC Andersen Flyttefirma Invoice'

def test_InsertDynamicData(
        templateFile: str,
        pdfName: str,
        contextFile: dict,
    ):
    context = contextFile

    templatePath = 'uploadedTemplate.docx'
    tpl = DocxTemplate(templateFile)
    
    context = InitializeContext(context)
    
    errMsg, valid = ValidateVariables(templatePath, context)
    if valid == False:
        return ValueError, errMsg

    # Inserts data from the context-dictionary to the template
    tpl.render(context)
    
    tpl.save(f'{pdfName}.docx')
    InsertPageNumbers(f'{pdfName}.docx')
    return ConvertDocxToPDF(pdfName)

class TestInsertDynamicData(unittest.TestCase):
    def test_contextTooLarge(self):
        testContext = deepcopy(contextDict)
        expectedErrMsg = ValueError, 'test1 was not found in template. '

        testContext['test1'] = 1
        self.assertEqual(test_InsertDynamicData(testTemplatePath, testPDFName, testContext), expectedErrMsg)
    
    def test_contextTooSmall(self):
        testContext = deepcopy(contextDict)
        expectedErrMsg = ValueError, 'Name was not found in context. '

        testContext.pop('Name')
        self.assertEqual(test_InsertDynamicData(testTemplatePath, testPDFName, testContext), expectedErrMsg)
    
    def test_success(self):
        self.assertTrue(test_InsertDynamicData(testTemplatePath, testPDFName, contextDict))


# ------------------------------------------------------------------------------------------- #


if __name__ == '__main__':
    unitTest = 0

    if unitTest == 1:
        unittest.main()
    else:
        import uvicorn
        uvicorn.run(app, host = '127.0.0.1', port = 8000)
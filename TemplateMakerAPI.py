# pip install 'fastapi[standard]'
# pip install python-docx
# pip install docx2pdf
# pip install python-dateutil

# Used in API calls
from fastapi import FastAPI, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse

# Used in MakeTemplate() and/or InsertDynamicData()
from docx import Document
import winreg
from docx.shared import Inches
from datetime import datetime

# Used in ConvertDocxToPDF()
import os
from docx2pdf import convert

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)


@app.post('/generate-template/')
async def generate_template(
        fileName : str = Form(...), 
        image : str = Form(None), 
        imageWidth : float = Form(None), 
        fileNameOverride : bool = Form(None)
    ):
    return GenerateTemplate(fileName, image, imageWidth, fileNameOverride, True)

@app.post('/insert-dynamic-data/')
async def inser_dynamic_data(
        fileName : str,
        items : dict,
        values : dict,
        totalPages : int
    ):
    return InsertDynamicData(fileName = fileName, values = values, items = items, totalPages = totalPages, postman = True)






def GenerateTemplate(fileName : str, image : str = None, imageWidth : str = None, fileNameOverride : bool = None, postman : bool = False):
    folderPath = FindFolderPath()
    if fileNameOverride != True:
        templatePath = f'{folderPath}/{fileName} Invoice TemplateFile'
    else:
        templatePath = f'{folderPath}/{fileName}'

    doc = Document()

    doc.add_paragraph('{{Page Number}} of {{Total Pages}}').alignment = 2

    doc.add_heading('Invoice', 0)

    if image != None:
        if imageWidth != None:
            doc.add_picture(image, width = Inches(imageWidth))
        else:
            doc.add_picture(image, width = Inches(1.5))
    
    doc.add_paragraph('Dear {{Name}} ,')

    p = doc.add_paragraph('Please find attached invoice for your recent purchase of ')
    p.add_run('\n\n{{Product(s)}}')

    table = doc.add_table(rows = 1, cols = 4)
    hdrCells = table.rows[0].cells
    hdrCells[0].text = 'Product Name'
    hdrCells[1].text = 'Units'
    hdrCells[2].text = 'Unit Price'
    hdrCells[3].text = 'Total Price'
    for i in range(4):
        hdrCells[i].paragraphs[0].runs[0].font.bold = True
    
    rowCells = table.add_row().cells
    rowCells[0].text = '{{Product Name}}'
    rowCells[1].text = '{{Units}}'
    rowCells[2].text = '{{Unit Price}}'
    rowCells[3].text = '{{Total Price}}'

    doc.add_page_break()

    doc.add_paragraph('{{Page Number}} of {{Total Pages}}').alignment = 2

    doc.add_paragraph('Date of task execution: {{Date of Execution}}')
    doc.add_paragraph('Please ensure that all dues are paid in full before {{Date of Payment}}')
    
    doc.add_picture('InvoEZ logo.png', width = Inches(4)).keep_together = True

    doc.add_paragraph('{{Comments}}')

    doc.add_paragraph('We apprecieate your business and please come again!').keep_together = True
    doc.add_paragraph('Sincerely').keep_together = True
    doc.add_paragraph('InvoEZ').keep_together = True

    doc.save(f'{templatePath}.docx')
    if postman == True: return FileResponse(f'{templatePath}.docx', media_type = 'application/docx', filename = fileName)





def InsertDynamicData(fileName : str, values : dict, items : dict, totalPages : int = None, postman : bool = False):
    folderPath = FindFolderPath()

    templatePath = f'{folderPath}/{fileName} Invoice TemplateFile'
    filePath = f'{folderPath}/{fileName} Invoice'
    tempPath = f'{folderPath}/{fileName} TempFile'

    doc = Document(f'{templatePath}.docx')

    valuesCount = 0
    for p in doc.paragraphs:
        if p.text.__contains__('{{'):
            if not p.text.__contains__('{{Page Number}} of {{Total Pages}}'):
                valuesCount += 1

    if len(values) != valuesCount:
        print(f'There are {len(values)} inputted values, but the template holds {valuesCount} values.')
        return
    
    for key in values:
        for p in doc.paragraphs:
            if p.text.__contains__(key):
                if key == 'Second Adresses' or 'Details':
                    p.text = values[key[0]]
                    for item in values[key]:
                        p.add_run(f'{item}\n')
                else:
                    p.text = values[key]

    table = None
    for t in doc.tables:
        if t.cell(0, 0).text.__contains__('{{Product Table}}'):
            t.cell(0, 0).text = ''
            table = t
            break

    if table != None:
        for key in items:
            rowCells = table.add_row().cells
            rowCells[0].text = key
            if items[key][0] > 1:
                if isinstance(items[key][0], float):
                    rowCells[1].text = f'{items[key][0]:,.2f}'
                elif isinstance(items[key][0], int):
                    rowCells[1].text = f'{items[key][0]}'
            rowCells[2].text = f'{items[key][1]:,.2f}'
            rowCells[3].text = f'{items[key][0] * items[key][1]:,.2f}'

        totalUnits = 0
        totalPrice = 0
        for key in items:
            totalUnits += items[key][0]
            totalPrice += items[key][0] * items[key][1]

        tableTotals = {'Total ex. moms' : totalUnits, 'Total ink. moms' : totalUnits}
        for key in tableTotals:
            rowCells = table.add_row().cells
            rowCells[0].text = key
            rowCells[1].text = f'{tableTotals[key]:,.2f}'
            rowCells[3].text = f'{totalPrice:,.2f}'
    
    doc.save(f'{tempPath}.docx')
    if totalPages == None: totalPages = CountPages(f'{tempPath}.docx')
    
    pageCount = 0
    for paragraph in doc.paragraphs:
        if paragraph.text.__contains__('{{Page Number}} of {{Total Pages}}'):
            pageCount += 1
            paragraph.text = f'Page {pageCount} of {totalPages}'
        
    doc.save(f'{tempPath}.docx')
    if postman == True: return ConvertDocxToPDF(filePath, tempPath, f'{fileName} Invoice')
    else: ConvertDocxToPDF(filePath, tempPath, f'{fileName} Invoice')




def CountPages(path):
    doc = Document(path)
    pages = 1    
    import re
    for p in doc.paragraphs:
        r = re.match('Chapter \d+', p.text)
        if r:
            print(r.group(), pages)
        for run in p.runs:
            if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                pages += 1

    return pages

def ConvertDocxToPDF(path, tempPath, fileName):
    convert(f'{tempPath}.docx', f'{path}.pdf')
    os.remove(f'{tempPath}.docx')

    return FileResponse(f'{path}.pdf', media_type = 'application/pdf', filename = fileName)

def FindFolderPath():
    reg_key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders')
    folderPath = winreg.QueryValueEx(reg_key, '{374DE290-123F-4565-9164-39C4925E467B}')[0]
    winreg.CloseKey(reg_key)

    return folderPath

def DeleteParagraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None




if __name__ == '__main__':
    postman = 0

    if postman == 0:
        items = {'Product 1' : [10, 99.95], 'Product 2' : [15, 199.95], 'Product 3' : [18, 324.95], 'Product 4' : [16, 499.95], 'Product 5' : [19, 649.95],
            'Product 6' : [4, 1499.95], 'Product 7' : [34, 124.95], 'Product 8' : [150, 1749.95], 'Product 9' : [6, 14999.95], 'Product 10' : [60, 19.95], 'Work Hours' : [22.5, 450]}
        fileName = 'Cadana Invoice TemplateFile'
        
        name = 'Cadana'
        customerName = 'Cadana Customer'

        image = 'CadanaLogo.png'
        imageWidth = 3

        '--------------------------------------------------'
        
        itemsHCFlyt = {'Fast pris for flytning' : [1, 9999.95], 'Fast pris for nedpakning' : [1, 249.95], 'Fast pris for udpakning' : [1, 499.95], 'Pris for opbevaring' : [1, 349.95],
            'Pris for leje af udstyr' : [8, 199.95], 'Tungløft' : [7, 99.95], 'Ekstra Arbejdstimer' : [7.5, 50]}
        valuesHCFlyt = {'Name' : 'Kenneth', 'Adress' : 'Kenneths address', 'City, postcode' : 'kenneths city and postcode', 'Customer Number' : 'Kundenummer:	123456', 'Order Number' : 'Ordrenummer:	123456', 
                        'Current Date' : f'Dato: {datetime.today().strftime("%d-%m-%Y")}', 'Offer Name' : 'Offer for Kenneth', 'Date of Execution' : 'Dato for flytning: 01-01-2025', 'Floor w Elevator' : 'Etage (med/uden elevator): St. til 1st floor, w Elevator', 
                        'Square Meters' : '5 m2', 'Parking' : 'Parkering og adgangsforhold:\nYes', 'Time Estimation' : 'Andet: 3 hours', 'Task Description' : 'Opgaver beskrivelse: This is the task description', 'Second Adresses' : ['Arbejdsadresse forskellig fra faktureringsadresse/Køreplan:', '1st location', '2nd location'], 
                        'Comments' : 'No comments', 'Agreed Date' : 'D. 01-01-2025 mellem kl. 10:00 og 11:00', 'Agreed Date for Equipment' : 'D. 02-01-2025 mellem kl. 11:00 og 12:00', 
                        'Customer Phone' : 'Telefon 12345678', 'Total Price' : 'Total pris for opgaven	10.000 kr', 'Reg Number' : 'Reg: 123456', 'Account Number' : 'Konto: 123456', 'Order Id' : 'Anfør 123456', 'Details' : ['This is a details', 'This is also a detail']}
        fileNameHCFlyt =  'Ordrebekræftelse_opdateret'

        customerNameHCFlyt = 'H.C. Andersens Flyttefirma A/S'
        commentsHCFlyt = ['This is a comment', 'This is another comment']
        paymentDay = '01-01-2025'

        '--------------------------------------------------'

        generateTemplate = 0

        if generateTemplate == 1:
            GenerateTemplate(fileName, image, imageWidth, True)
        else:
            InsertDynamicData(name, valuesHCFlyt, itemsHCFlyt, 8)
    else:
        import uvicorn
        uvicorn.run(app, host='127.0.0.1', port=8000)
from fastapi import FastAPI, File, UploadFile, Form
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import os

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)

@app.post("/generate-pdf/")
async def generate_pdf(
    wordTemplate: UploadFile = File(...),
    company_name: str = Form(...),
    date: str = Form(...),
    hours_of_work: str = Form(...),
    hour_price: str = Form(...)
):
    # Saving the uploaded word file/template
    template_path = "template.docx"
    with open(template_path, "wb") as buffer:
        buffer.write(wordTemplate.file.read())

    # Loading template
    doc = Document(template_path)

    # Replacement of placeholders in template
    for paragraph in doc.paragraphs:
        if "{{company_name}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{company_name}}", company_name)
        if "{{date}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{date}}", date)
        if "{{hours_of_work}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{hours_of_work}}", hours_of_work)
        if "{{hour_price}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{hour_price}}", hour_price)

    # Saving the modified document
    modified_doc_path = "modified_template.docx"
    doc.save(modified_doc_path)

    # Converting the modified document to a PDF
    pdf_path = "output.pdf"
    pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    for paragraph in doc.paragraphs:
        story.append(Paragraph(paragraph.text, styles["Normal"]))

    pdf_doc.build(story)

    # Removal of temporary files
    os.remove(template_path)
    os.remove(modified_doc_path)

    return FileResponse(pdf_path, media_type='application/pdf', filename="output.pdf")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
from fastapi import FastAPI, File, UploadFile, Form
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Image, Table, TableStyle, Flowable
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib import colors
import os


app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)

# @app.get("/")
# def index():
#     return {"message": "PDF Generator API"}


#pdf_data = {}

@app.get("/")
async def index():
    # Define the path for the generated PDF
    pdf_path = "output.pdf"
    
    # Generate a simple PDF file with static content
    pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story: list[Flowable] = [Paragraph("This is the content of the generated PDF file.", styles["Normal"])]
    
    # Build the PDF
    pdf_doc.build(story)

    # Return the PDF inline
    return FileResponse(
        pdf_path,
        media_type="application/pdf",
        headers={"Content-Disposition": "inline; filename=output.pdf"}
    )


############################################################################################################

@app.post("/generate-pdf/")
async def generate_pdf(
    wordTemplate: UploadFile = File(...),
    company_name: str = Form(...),
    date: str = Form(...),
    hours_of_work: str = Form(...),
    hour_price: str = Form(...),
    images: list[UploadFile] = File(None)  
):
    #global pdf_data

    #Saving the uploaded files
    template_path = "template.docx"
    with open(template_path, "wb") as buffer:
        buffer.write(wordTemplate.file.read())

    image_paths = []
    if images:
        for i, image in enumerate(images):
            image_path = f"image_{i}.png"
            with open(image_path, "wb") as buffer:
                buffer.write(image.file.read())
            image_paths.append(image_path)

    #Loading the template
    doc = Document(template_path)

    #Replacing the placeholders
    for paragraph in doc.paragraphs:
        if "{{company_name}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{company_name}}", company_name)
        if "{{date}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{date}}", date)
        if "{{hours_of_work}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{hours_of_work}}", hours_of_work)
        if "{{hour_price}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{hour_price}}", hour_price)

    #Saving the modified template
    modified_doc_path = "modified_template.docx"
    doc.save(modified_doc_path)

    #Converting the modified template to PDF
    pdf_path = "output.pdf"
    pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story: list[Flowable] = []

    image_index = 0
    for paragraph in doc.paragraphs:
        if "{{image}}" in paragraph.text:
            #Adding any text before the placeholder
            parts = paragraph.text.split("{{image}}")
            if parts[0]:
                story.append(Paragraph(parts[0], styles["Normal"]))
            #Adding the images in a 4-column layout
            if image_paths:
                table_data = []
                row = []
                for image_path in image_paths:
                    img = Image(image_path)
                    max_width = 1.5 * inch  #Maximum width for the image
                    max_height = 1.5 * inch  #Maximum height for the image

                    #Resizing the images to keep aspect ratio
                    if img.drawWidth > max_width or img.drawHeight > max_height:
                        aspect_ratio = img.drawWidth / img.drawHeight
                        if img.drawWidth > img.drawHeight:
                            img.drawWidth = max_width
                            img.drawHeight = max_width / aspect_ratio
                        else:
                            img.drawHeight = max_height
                            img.drawWidth = max_height * aspect_ratio

                    row.append(img)
                    if len(row) == 4:
                        table_data.append(row)
                        row = []
                if row:
                    table_data.append(row)

                table = Table(table_data)
                table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),
                    ('BOX', (0, 0), (-1, -1), 0.25, colors.black),
                ]))
                story.append(table)
                image_index += len(image_paths)

            #Adding text after the placeholder
            if len(parts) > 1 and parts[1]:
                story.append(Paragraph(parts[1], styles["Normal"]))
        else:
            story.append(Paragraph(paragraph.text, styles["Normal"]))

    pdf_doc.build(story)

    #Removal of temporary files
    os.remove(template_path)
    os.remove(modified_doc_path)
    for image_path in image_paths:
        os.remove(image_path)

    return FileResponse(pdf_path, 
    media_type='application/pdf', 
    headers={"Content-Disposition": "inline; filename=output.pdf"})
    


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
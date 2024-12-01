from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Image, Table
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
import os

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)

@app.post("/generate-pdf/")
#list of form fields. 
async def generate_pdf(
    wordTemplate: UploadFile = File(...),
    company_name: str = Form(None),
    date: str = Form(None),
    hours_of_work: str = Form(None),
    hour_price: str = Form(None),
    images: list[UploadFile] = File(None),
    full_name: str = Form(...),
    address: str = Form(...),
    city: str = Form(...),
    zip_code: str = Form(...),
    customer_number: str = Form(...),
    order_number: str = Form(...),
    order_type: str = Form(...),
    offer_name: str = Form(...),
    move_date: str = Form(...),
    elevator: str = Form(...),
    job_sq_m: str = Form(...),
    parking_and_access: str = Form(...),
    other_job_info: str = Form(...),
    project_description: str = Form(...),
    timetable: str = Form(...),
    comments: str = Form(...),
    pickup_date: str = Form(...),
    pickup_starttime: str = Form(...),
    pickup_endtime: str = Form(...),
    delivery_date: str = Form(...),
    delivery_starttime: str = Form(...),
    delivery_endtime: str = Form(...),
    customer_phone: str = Form(...),
    total_price_incl: str = Form(...),
    reg_number: str = Form(...),
    account_number: str = Form(...),
    storage_sq_m: str = Form(...),
    rental_amount: str = Form(...),
    heavylift_amount: str = Form(...),
    rental_price: str = Form(...),
    heavylift_price: str = Form(...),
    moving_price_total: str = Form(...),
    packing_price_total: str = Form(...),
    unpacking_price_total: str = Form(...),
    storage_price_total: str = Form(...),
    rental_price_total: str = Form(...),
    heavylift_price_total: str = Form(...),
    total_price_ex: str = Form(...),
    standard_product_form: str = Form(...),
    prices: str = Form(...),
    conditions: str = Form(...)
):
    #Saving the template
    template_path = "template.docx"
    with open(template_path, "wb") as buffer:
        buffer.write(wordTemplate.file.read())

    #Saving images
    image_paths = []
    if images:
        for i, image in enumerate(images):
            image_path = f"image_{i}.png"
            with open(image_path, "wb") as buffer:
                buffer.write(image.file.read())
            image_paths.append(image_path)

    #Dictionary of placeholders and values
    placeholders = {
        "{{full_name}}": full_name,
        "{{address}}": address,
        "{{city}}": city,
        "{{zip_code}}": zip_code,
        "{{customer_number}}": customer_number,
        "{{order_number}}": order_number,
        "{{order_type}}": order_type,
        "{{offer_name}}": offer_name,
        "{{move_date}}": move_date,
        "{{elevator}}": elevator,
        "{{job_sq_m}}": job_sq_m,
        "{{parking_and_access}}": parking_and_access,
        "{{other_job_info}}": other_job_info,
        "{{project_description}}": project_description,
        "{{timetable}}": timetable,
        "{{comments}}": comments,
        "{{pickup_date}}": pickup_date,
        "{{pickup_starttime}}": pickup_starttime,
        "{{pickup_endtime}}": pickup_endtime,
        "{{delivery_date}}": delivery_date,
        "{{delivery_starttime}}": delivery_starttime,
        "{{delivery_endtime}}": delivery_endtime,
        "{{customer_phone}}": customer_phone,
        "{{total_price_incl}}": total_price_incl,
        "{{reg_number}}": reg_number,
        "{{account_number}}": account_number,
        "{{storage_sq_m}}": storage_sq_m,
        "{{rental_amount}}": rental_amount,
        "{{heavylift_amount}}": heavylift_amount,
        "{{rental_price}}": rental_price,
        "{{heavylift_price}}": heavylift_price,
        "{{moving_price_total}}": moving_price_total,
        "{{packing_price_total}}": packing_price_total,
        "{{unpacking_price_total}}": unpacking_price_total,
        "{{storage_price_total}}": storage_price_total,
        "{{rental_price_total}}": rental_price_total,
        "{{heavylift_price_total}}": heavylift_price_total,
        "{{total_price_ex}}": total_price_ex,
        "{{standard_product_form}}": standard_product_form,
        "{{prices}}": prices,
        "{{conditions}}": conditions,
        "{{date}}": date,
        "{{hours_of_work}}": hours_of_work,
        "{{hour_price}}": hour_price
    }

    #Replacing the placeholders
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text and value is not None:
                paragraph.text = paragraph.text.replace(placeholder, value)

    #Saving the modified template
    modified_doc_path = "modified_template.docx"
    doc.save(modified_doc_path)

    #Converting the modified template to PDF
    pdf_path = "output.pdf"
    pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    #Adding images to the PDF
    image_index = 0
    for paragraph in doc.paragraphs:
        if "{{image}}" in paragraph.text:
            parts = paragraph.text.split("{{image}}")
            if parts[0]:
                story.append(Paragraph(parts[0], styles["Normal"]))
            if image_paths:
                table_data = []
                row = []
                #rezising images to keep aspect ratio
                for image_path in image_paths:
                    img = Image(image_path)
                    max_width = 1.5 * inch  
                    max_height = 1.5 * inch  
                    img.drawWidth = max_width
                    img.drawHeight = max_height * img.imageHeight / img.imageWidth
                    if img.drawHeight > max_height:
                        img.drawHeight = max_height
                        img.drawWidth = max_height * img.imageWidth / img.imageHeight
                    if img.drawWidth > max_width:
                        img.drawWidth = max_width
                        img.drawHeight = max_width * img.imageHeight / img.imageWidth
                    #puttin them into a 4-column layout
                    row.append(img)
                    if len(row) == 4:
                        table_data.append(row)
                        row = []
                if row:
                    table_data.append(row)
                story.append(Table(table_data, colWidths=[1.5 * inch] * 4))
        else:
            story.append(Paragraph(paragraph.text, styles["Normal"]))

    
    pdf_doc.build(story)

    #Removing tempoary files
    os.remove(template_path)
    os.remove(modified_doc_path)
    for image_path in image_paths:
        os.remove(image_path)


    return FileResponse(pdf_path, media_type='application/pdf', filename="output.pdf")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
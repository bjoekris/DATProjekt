# Used for inserting dynamic data to word template
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Mm
from fastapi.responses import FileResponse
from PIL import Image

# Used for converting from docx to pdf
import os
import docx2pdf

# Used for retrieving image files from URL
import requests

def InsertDynamicData(
        templatePath: str,
        context: dict,
        pdfName: str,
    ):
    # Finds and open the copied template as a DocxTemplate object
    tpl = DocxTemplate(templatePath)
    
    # Finds images from URL and puts them into the context-dictionary
    if context.__contains__('Images'):
        HandleImages(context, tpl)

    # Validates the context- and template-variables
    errMsg, valid = ValidateVariables(templatePath, context)
    if valid == False:
        raise ValueError(errMsg)

    # Inserts data from the context-dictionary to the template
    tpl.render(context)
    
    # Saves and converts to PDF
    tpl.save(f'{pdfName}.docx')
    
    # Returns final PDF
    return ConvertDocxToPDF(pdfName, templatePath)
    
def ConvertDocxToPDF(path: str, templatePath: str):
    docx2pdf.convert(f'{path}.docx', f'{path}.pdf')
    RemoveTempFiles(templatePath, path)

    return FileResponse(f'{path}.pdf', media_type = 'application/pdf', filename = path)

def ValidateVariables(path: str, context: dict):
    doc = Document(path)
    valid = True

    keysContained = []
    keysNotContained = []
    values = []
    valuesNotInputted = []
    valuesInputted = []

    # All variables in the template must start with "{{"
    # We can therefore identify the variables, by searching the document for these
    for p in doc.paragraphs:
        if p.text.__contains__('{{') and not (p.text.__contains__('{% for') or p.text.__contains__('{% endfor')):
            tempValues = FindVariables(p)
            for value in tempValues:
                if not values.__contains__(value):
                    values.append(value)
    
    for t in doc.tables:
        tableEligible = True
        for c in t._cells:
            if tableEligible == True:
                if c.text.__contains__('{%tr') or c.text.__contains__('{%tc'):
                    tableEligible = False

                else:
                    for p in c.paragraphs:
                        if p.text.__contains__('{{') and not (p.text.__contains__('{% for') or p.text.__contains__('{% endfor')):
                            tempValues = FindVariables(p)
                            for value in tempValues:
                                if not values.__contains__(value):
                                    values.append(value)
            
            else:
                if c.text.__contains__('{%tr endfor') or c.text.__contains__('{%tc endfor'):
                    tableEligible = True

    # Ignores lists, since the are allowed to be empty
    for key in context:
        if not isinstance(context[key], list) and not isinstance(context[key], dict): keysNotContained.append(key)
    
    for value in values:
        valuesNotInputted.append(value)
    
    for value in values:
        for key in context:
            if value == key:
                keysContained.append(key)
                keysNotContained.remove(key)

                valuesInputted.append(value)
                valuesNotInputted.remove(value)
    
    # Builds the exception string, so it can be returned to the user
    errorMsg = ''
    if not len(keysNotContained) == 0:
        for key in keysNotContained:
            errorMsg += f'{key} was not found in template. '
        
        valid = False
    
    if not len(valuesNotInputted) == 0:
        for value in valuesNotInputted:
            errorMsg += f'{value} was not found in context. '
        
        valid = False
    
    return errorMsg, valid

def FindVariables(t):
    tempValues = t.text.rsplit('{{')

    for temp in tempValues:
        if not temp.__contains__('}}'):
            tempValues.remove(temp)
    
    # Isolates the template varaiables from the final clamps, so that they can be compaired with the context variables
    foundValues = []
    for i in range(len(tempValues)):
        tempValue = tempValues[i].split('}}')[0]
        foundValues.append(tempValue)

    return foundValues

def FindImage(url: str, fileName: str):
    data = requests.get(url).content
    path = open(f'{fileName}.png','wb')
    path.write(data)
    path.close()

    return f'{fileName}.png'

def HandleImages(context: dict, tpl: DocxTemplate):
    index = 0
    
    for image in context['Images']:
        foundImage = FindImage(image['URL'], f'image{index}')
        templateImage = InlineImage(tpl, foundImage)

        maxWidth = 170
        maxHeight = 125

        img = Image.open(foundImage)
        imageRatio = img.width / img.height

        # This is where it is determined whether to use, Size, Width, or Height from the given image values
        # This is also where it is checked if a size-value has been set too high
        if image['Size'] > image['Width'] and image['Size'] > image['Height'] and image['Size'] > 0:
            if img.width > img.height: useWidth = True
            else: useWidth = False
            size = image['Size']
        
        elif image['Width'] > 0 or image['Height'] > 0:
            if image['Width'] > image['Height']: useWidth = True
            else: useWidth = False
            
            if useWidth == True: size = image['Width']
            else: size = image['Height']
        
        else:
            img.close()
            raise ValueError(f'{image["URL"]} caused an error, at position: {index + 1}. At least one between "Size", "Width", and "Height" must be larger than 0.')

        if useWidth == True:
            templateImage = InlineImage(tpl, foundImage, width = Mm(size))
            newHeight = size / imageRatio

            if templateImage.width > Mm(maxWidth):
                img.close()
                raise ValueError(f'{image["URL"]} with Width: {size}, at position: {index + 1} has exceeded the maximum width og 170.')
            
            elif newHeight > maxHeight:
                img.close()
                raise ValueError(f'{image["URL"]} with Height: {newHeight}, at position: {index + 1} has exceeded the maximum height og 125.')
        
        else:
            templateImage = InlineImage(tpl, foundImage, height = Mm(size))
            newWidth = imageRatio * size

            if templateImage.height > Mm(maxHeight):
                img.close()
                raise ValueError(f'{image["URL"]} with Height: {size}, at position: {index + 1} has exceeded the maximum height og 125.')

            if newWidth > maxWidth:
                img.close()
                raise ValueError(f'{image["URL"]} with Width: {newWidth}, at position: {index + 1} has exceeded the maximum width og 170.')

        img.close()

        if image['Positioned'] == 'True':
            context[f'Image{index}'] = templateImage
        
        else:
            if not context.__contains__(f'Images{image["List"]}'):
                context[f'Images{image["List"]}'] = []
            
            context[f'Images{image["List"]}'].append(templateImage)

        index += 1
        
    context.pop('Images')

def RemoveTempFiles(templatePath: str, tempDocxPath: str):
    if os.path.isfile(templatePath): os.remove(templatePath)
    if os.path.isfile(f'{tempDocxPath}.docx'): os.remove(f'{tempDocxPath}.docx')
    
    imagesToRemove = FindAllImages()
    if len(imagesToRemove) > 0:
        for image in imagesToRemove: os.remove(image)

def FindAllImages():
    imagesFound = []
    for f in os.listdir():
        if f.__contains__('.png'): imagesFound.append(f)
    
    return imagesFound
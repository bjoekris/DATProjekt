# Brugt til at indsætte dynamiske data i word skabelon
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Mm
from fastapi.responses import FileResponse
from PIL import Image

# Brugt til at konvertere fra docx-format til pdf-format
import os
import docx2pdf

# Brugt til at finde et givent billede, fra URL
import requests

def InsertDynamicData(
        templatePath: str,
        context: dict,
        pdfName: str,
    ):
    ## -------------------------------------- Skrevet af Magnus -------------------------------------- ##
    # Finder og åbner den kopierede skabelon som DocxTemplate objekt
    tpl = DocxTemplate(templatePath)
    
    # Finder givene billeder fra URL, og ligger den i context-dictionary
    if context.__contains__('Images'):
        HandleImages(context, tpl)
    
    # Validere context- og skabelon-variabler
    errMsg, valid = ValidateVariables(templatePath, context)
    if valid == False:
        raise ValueError(errMsg)

    # Indsætter data fra context-dictionary til skabelon
    tpl.render(context)
    
    # Gemmer og konvertere til PDF
    tpl.save(f'{pdfName}.docx')
    
    # Returnere færdig PDF
    return ConvertDocxToPDF(pdfName, templatePath)
    ## ---------------------------------------------------------------------------------------------- ##

## ---------------------------------------- Skrevet af Magnus --------------------------------------- ##
def ConvertDocxToPDF(path: str, templatePath: str):
    docx2pdf.convert(f'{path}.docx', f'{path}.pdf')
    RemoveTempFiles(templatePath, path)
    ## ---------------------------------------------------------------------------------------------- ##

    ## ------------------------------------ Skrevet af Bjørn ---------------------------------------- ##
    return FileResponse(f'{path}.pdf', media_type = 'application/pdf', filename = path)
    ## ---------------------------------------------------------------------------------------------- ##

## ---------------------------------------- Skrevet af Magnus --------------------------------------- ##
def ValidateVariables(path: str, context: dict):
    doc = Document(path)
    valid = True

    keysContained = []
    keysNotContained = []
    values = []
    valuesNotInputted = []
    valuesInputted = []

    # Alle variabler i skabelonet starter med "{{"
    # Derfor kan man finde alle variabler ved at lede gennem alt tekst i dokumentet efter disse
    for p in doc.paragraphs:
        if p.text.__contains__('{{') and not (p.text.__contains__('{% for') or p.text.__contains__('{% endfor')):
            tempValues = FindVariables(p)
            for value in tempValues:
                if not values.__contains__(value):
                    values.append(value)
    
    for t in doc.tables:
        tableEligible = True
        for c in t._cells:
            if tableEligible == True:
                if c.text.__contains__('{%tr') or c.text.__contains__('{%tc'):
                    tableEligible = False

                else:
                    for p in c.paragraphs:
                        if p.text.__contains__('{{') and not (p.text.__contains__('{% for') or p.text.__contains__('{% endfor')):
                            tempValues = FindVariables(p)
                            for value in tempValues:
                                if not values.__contains__(value):
                                    values.append(value)
            
            else:
                if c.text.__contains__('{%tr endfor') or c.text.__contains__('{%tc endfor'):
                    tableEligible = True

    # Ignoere lister, da disse godt må være tomme
    for key in context:
        if not isinstance(context[key], list): keysNotContained.append(key)
    
    for value in values:
        valuesNotInputted.append(value)
    
    for value in values:
        for key in context:
            if value == key:
                keysContained.append(key)
                keysNotContained.remove(key)

                valuesInputted.append(value)
                valuesNotInputted.remove(value)
    
    # Bygger exception-string, så den kan returneres til brugeren
    errMsg = ''
    if not len(keysNotContained) == 0:
        for key in keysNotContained:
            errMsg += f'{key} was not found in template. '
        
        valid = False
    
    if not len(valuesNotInputted) == 0:
        for value in valuesNotInputted:
            errMsg += f'{value} was not found in context. '
        
        valid = False
    
    return errMsg, valid

def FindVariables(t):
    tempValues = t.text.rsplit('{{')

    for temp in tempValues:
        if not temp.__contains__('}}'):
            tempValues.remove(temp)

    # Isolere skabelon-variablerne fra de sidste klammer, så de kan sammenlignes med context-variablerne
    foundValues = []
    for i in range(len(tempValues)):
        tempValue = tempValues[i].split('}}')[0]
        foundValues.append(tempValue)

    return foundValues

def FindImage(url: str, fileName: str):
    data = requests.get(url).content
    f = open(f'{fileName}.png',"wb")
    f.write(data)
    f.close()

    return f'{fileName}.png'

def HandleImages(context: dict, tpl: DocxTemplate):
    index = 0
    
    for image in context['Images']:
        foundImage = FindImage(image['URL'], f'image{index}')
        templateImage = InlineImage(tpl, foundImage)

        maxWidth = 170
        maxHeight = 125

        img = Image.open(foundImage)
        imageRatio = img.width / img.height

        # Her bestemes der om hvilken af størrelses-værdierne der skal bruges fra det givne billede
        # Det er også her der bliver checked om størrelses-værdierne er sat for højt eller lavt
        if image['Size'] == 0:
            img.close()
            raise ValueError(f'{image["URL"]} caused an error, at position: {index + 1}. Size must be larger than 0.')
        
        if image['Option'] == 'Auto':
            if img.width > img.height: useWidth = True
            else: useWidth = False
        
        else:
            if image['Option'] == 'Width': useWidth = True
            else: useWidth = False
        
        if useWidth == True:
            templateImage = InlineImage(tpl, foundImage, width = Mm(image['Size']))
            newHeight = image['Size'] / imageRatio

            if templateImage.width > Mm(maxWidth):
                img.close()
                raise ValueError(f'{image["URL"]} with Width: {image["Size"]}, at position: {index + 1} has exceeded the maximum width og 170.')
            
            elif newHeight > maxHeight:
                img.close()
                raise ValueError(f'{image["URL"]} with Height: {newHeight}, at position: {index + 1} has exceeded the maximum height og 125.')
        
        else:
            templateImage = InlineImage(tpl, foundImage, height = Mm(image["Size"]))
            newWidth = imageRatio * image['Size']

            if templateImage.height > Mm(maxHeight):
                img.close()
                raise ValueError(f'{image["URL"]} with Height: {image["Size"]}, at position: {index + 1} has exceeded the maximum height og 125.')

            if newWidth > maxWidth:
                img.close()
                raise ValueError(f'{image["URL"]} with Width: {newWidth}, at position: {index + 1} has exceeded the maximum width og 170.')

        img.close()

        if image['List'] == 0:
            context[f'Image{index}'] = templateImage
        
        else:
            if not context.__contains__(f'Images{image["List"]}'):
                context[f'Images{image["List"]}'] = []
            
            context[f'Images{image["List"]}'].append(templateImage)

        index += 1
        
    context.pop('Images')

def RemoveTempFiles(templatePath: str, tempDocxPath: str):
    if os.path.isfile(templatePath): os.remove(templatePath)
    if os.path.isfile(f'{tempDocxPath}.docx'): os.remove(f'{tempDocxPath}.docx')
    
    imagesToRemove = FindAllImages()
    if len(imagesToRemove) > 0:
        for image in imagesToRemove: os.remove(image)

def FindAllImages():
    imagesFound = []
    for f in os.listdir():
        if f.__contains__('.png'): imagesFound.append(f)
    
    return imagesFound
## -------------------------------------------------------------------------------------------------- ##